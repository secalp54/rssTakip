import xml.etree.ElementTree as ET
from docx import Document
import requests
import feedparser

input("Rss programına hoşgeldiniz.Devam etmek için enter tuşuna basınız.")
items=[]
class Item:
  def __init__(self, link, title,published,description):
    self.link = link
    self.title = title
    self.published = published
    self.description = description

f = open("haber.txt", "r")

for line in f:
    line=line[:-1]
    
    root=feedparser.parse(line)
    for item in root.entries:
        link = item.link
        title = item.title
        published= item.published
        description=item.description
        items.append(Item(link,title,published,description))
       

print(len(items))
bulunanLinkler=[]
bulunanTitle=[]
bulunanDesc=[]
kelimeler=[]
keywords= open("keywords.txt", "r", encoding="utf-8")
for kelime in keywords:
   kelimeler.append(kelime)
for haber in items:
   for k in range(0, len(kelimeler)):
      
      metin=haber.description.upper()
      aranan=kelimeler[k].upper()[:-1]
      bulduMu=metin.find(aranan)



      if bulduMu!=-1:
         bulunanLinkler.append(haber.link)
         bulunanTitle.append(haber.title+"-"+haber.published)
         bulunanDesc.append(haber.description)
         break

print(bulunanLinkler)
document = Document()
i=0
for bulunan in bulunanLinkler:
   print(bulunan)
   document.add_heading(bulunanTitle[i], level=1)
   document.add_paragraph(bulunan)
   document.add_paragraph(bulunanDesc[i])
   i=i+1
input("işlem tamamlandı. Linkler word dosyasın oluşturulacak. Devam için enter tuşuna basınız.")
document.save("linkler.docx")

